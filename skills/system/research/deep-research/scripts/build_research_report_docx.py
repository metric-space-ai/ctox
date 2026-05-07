#!/usr/bin/env python3
"""Build a Word research report from a structured JSON synthesis.

The script is intentionally small and schema-tolerant so agents can generate a
report JSON after deep research and obtain a consistent DOCX without hand-coding
Word layout every time.
"""

from __future__ import annotations

import argparse
import json
import tempfile
import urllib.request
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    payload = json.loads(args.input.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    add_title(doc, payload)
    add_metadata(doc, payload.get("metadata", {}))

    for figure in payload.get("figures", []):
        add_figure(doc, figure)

    for section in payload.get("sections", []):
        add_section(doc, section)

    for table in payload.get("tables", []):
        add_table(doc, table)

    references = payload.get("references", [])
    if references:
        doc.add_heading("Quellen", level=1)
        for index, ref in enumerate(references, start=1):
            add_reference(doc, index, ref)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)


def add_title(doc: Document, payload: dict[str, Any]) -> None:
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(payload.get("title", "Research Report"))
    run.bold = True
    run.font.color.rgb = RGBColor(30, 50, 75)

    subtitle = payload.get("subtitle")
    if subtitle:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(subtitle))
        run.font.size = Pt(12)
        run.italic = True


def add_metadata(doc: Document, metadata: dict[str, Any]) -> None:
    if not metadata:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in metadata.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = render_inline(value)
    doc.add_paragraph()


def add_figure(doc: Document, figure: dict[str, Any]) -> None:
    path = figure.get("path")
    image_path = Path(path) if path else None
    if image_path is None and figure.get("url"):
        image_path = download_figure(str(figure["url"]))
    if image_path and image_path.is_file():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Cm(float(figure.get("width_cm", 15))))
    caption = figure.get("caption")
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(caption))
        run.italic = True
        run.font.size = Pt(9)


def download_figure(url: str) -> Path | None:
    suffix = Path(url.split("?", 1)[0]).suffix.lower()
    if suffix not in {".png", ".jpg", ".jpeg", ".webp"}:
        suffix = ".img"
    target = Path(tempfile.gettempdir()) / f"ctox-research-figure-{abs(hash(url))}{suffix}"
    try:
        request = urllib.request.Request(url, headers={"User-Agent": "ctox-deep-research/0.1"})
        with urllib.request.urlopen(request, timeout=12) as response:
            target.write_bytes(response.read(5_000_000))
        return target
    except Exception:
        return None


def add_section(doc: Document, section: dict[str, Any]) -> None:
    doc.add_heading(str(section.get("heading", "Abschnitt")), level=int(section.get("level", 1)))
    for item in section.get("content", []):
        if isinstance(item, str):
            doc.add_paragraph(item)
        elif isinstance(item, dict):
            kind = item.get("type", "paragraph")
            if kind == "bullet":
                for value in item.get("items", []):
                    doc.add_paragraph(render_inline(value), style="List Bullet")
            elif kind == "numbered":
                for value in item.get("items", []):
                    doc.add_paragraph(render_inline(value), style="List Number")
            elif kind == "callout":
                p = doc.add_paragraph()
                run = p.add_run(render_inline(item.get("text", "")))
                run.bold = True
                run.font.color.rgb = RGBColor(125, 65, 0)
            else:
                doc.add_paragraph(render_inline(item.get("text", "")))


def add_table(doc: Document, spec: dict[str, Any]) -> None:
    title = spec.get("title")
    if title:
        doc.add_heading(str(title), level=int(spec.get("level", 2)))
    columns = [str(value) for value in spec.get("columns", [])]
    rows = spec.get("rows", [])
    if not columns:
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, column in enumerate(columns):
        cell = table.rows[0].cells[idx]
        cell.text = column
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for raw_row in rows:
        cells = table.add_row().cells
        values = raw_row if isinstance(raw_row, list) else [raw_row.get(col, "") for col in columns]
        for idx, value in enumerate(values[: len(columns)]):
            cells[idx].text = render_inline(value)
    doc.add_paragraph()


def add_reference(doc: Document, index: int, ref: Any) -> None:
    p = doc.add_paragraph()
    if isinstance(ref, dict):
        title = ref.get("title", "Quelle")
        url = ref.get("url", "")
        note = ref.get("note", "")
        text = f"[{index}] {title}"
        if note:
            text += f". {note}"
        if url:
            text += f". {url}"
    else:
        text = f"[{index}] {ref}"
    p.add_run(text)


def render_inline(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return "" if value is None else str(value)


if __name__ == "__main__":
    main()
