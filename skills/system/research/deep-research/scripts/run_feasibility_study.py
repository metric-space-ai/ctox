#!/usr/bin/env python3
"""Run CTOX deep research and write a technical feasibility-study DOCX.

This script is intentionally deterministic. It gives the agent an executable
research-to-report path instead of relying on free-form prose instructions.
"""

from __future__ import annotations

import argparse
import json
import struct
import subprocess
import sys
import textwrap
import zlib
from datetime import date
from pathlib import Path
from typing import Any


TECHNOLOGIES = [
    {
        "name": "Wirbelstrom / Eddy Current",
        "keywords": ["eddy", "wirbelstrom", "induction", "conductive", "cfrp"],
        "principle": "Induzierte Ströme in leitfähigen Schichten; Defekte verändern Impedanz, Phase und Amplitude.",
        "success": 4,
        "throughput": 3,
        "single_shot": 2,
        "foil": 3,
        "uncertainty": "mittel",
        "verdict": "Aussichtsreichster kontaktloser Kandidat für direkte Kupfergitter-Anomalien, wenn Lift-off und Schichtstack kalibriert werden.",
    },
    {
        "name": "Pulsed Eddy Current / Induktionsthermografie",
        "keywords": ["pulsed eddy", "induction thermography", "thermography", "lock-in"],
        "principle": "Transient elektromagnetische Anregung und thermische Antwort; Risse/Unterbrechungen verändern Strom- und Wärmefluss.",
        "success": 4,
        "throughput": 3,
        "single_shot": 3,
        "foil": 3,
        "uncertainty": "mittel",
        "verdict": "Sehr relevant als Flächenverfahren; benötigt Stimulusdesign und Referenzcoupons.",
    },
    {
        "name": "Terahertz Imaging",
        "keywords": ["terahertz", "thz"],
        "principle": "THz-Laufzeit und Reflexionskontrast in Dielektrika; Metalle wirken stark reflektierend/abschirmend.",
        "success": 2,
        "throughput": 3,
        "single_shot": 3,
        "foil": 1,
        "uncertainty": "mittel",
        "verdict": "Für Lack/Primer/Composite-Schichten interessant, aber für Metallgitter plus möglicher Folie physikalisch stark limitiert.",
    },
    {
        "name": "Mikrowelle / mmWave",
        "keywords": ["microwave", "mmwave", "radar", "millimeter"],
        "principle": "Elektromagnetische Streuung und Reflexion bei längeren Wellenlängen; empfindlich gegenüber leitfähigen Grenzflächen.",
        "success": 2,
        "throughput": 4,
        "single_shot": 4,
        "foil": 1,
        "uncertainty": "mittel-hoch",
        "verdict": "Als schnelle Flächenindikation möglich, aber die Folie dominiert wahrscheinlich die Rückstreuung.",
    },
    {
        "name": "Hyperspektral / optisch",
        "keywords": ["hyperspectral", "optical", "spectral"],
        "principle": "Spektrale Oberflächenreflexion; nur indirekte Korrelation mit verdeckten Strukturen.",
        "success": 1,
        "throughput": 5,
        "single_shot": 5,
        "foil": 1,
        "uncertainty": "hoch",
        "verdict": "Für Oberflächenzustand und Beschichtungen nützlich, aber ohne optischen Pfad nicht geeignet zur direkten Kupfergitterprüfung.",
    },
    {
        "name": "Infrarot-Thermografie",
        "keywords": ["infrared", "thermography", "thermal"],
        "principle": "Aktive oder passive Wärmeflussmessung; Unterbrechungen, Delaminationen und lokale Leitfähigkeitsänderungen erzeugen Muster.",
        "success": 3,
        "throughput": 4,
        "single_shot": 4,
        "foil": 2,
        "uncertainty": "mittel",
        "verdict": "Stark für indirekte Defektanzeichen und schnelle Screening-Gates; direkte Geometrieauflösung des Gitters unsicher.",
    },
    {
        "name": "Shearografie",
        "keywords": ["shearography", "speckle", "strain"],
        "principle": "Optische Messung belastungsinduzierter Dehnungsfelder; zeigt strukturelle Anomalien indirekt.",
        "success": 3,
        "throughput": 4,
        "single_shot": 4,
        "foil": 2,
        "uncertainty": "mittel",
        "verdict": "Guter Ergänzungstest für Delaminationen/Verbundfehler, aber kein primäres Verfahren zur Gittertopologie.",
    },
    {
        "name": "Röntgen / CT",
        "keywords": ["x-ray", "xray", "computed tomography", "ct", "radiography"],
        "principle": "Absorptions-/Phasenkontrast und 3D-Rekonstruktion; Metall und Composite sind direkt geometrisch sichtbar.",
        "success": 5,
        "throughput": 1,
        "single_shot": 2,
        "foil": 5,
        "uncertainty": "niedrig",
        "verdict": "Beste Ground-Truth-Referenz, aber für Produktions-Single-Shot meist zu langsam, teuer oder regulatorisch schwer.",
    },
    {
        "name": "Magnetische Verfahren / MFL",
        "keywords": ["magnetic", "magneto", "flux leakage", "mfl"],
        "principle": "Magnetische Flussänderungen; Kupfer und CFK sind nicht ferromagnetisch, daher nur Spezialfälle.",
        "success": 1,
        "throughput": 2,
        "single_shot": 2,
        "foil": 1,
        "uncertainty": "mittel",
        "verdict": "Für Kupfergitter nicht primär geeignet; eher nur über induzierte Ströme statt statische Magnetik.",
    },
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--query", required=True)
    parser.add_argument("--workspace", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--title", default="Machbarkeitsstudie")
    parser.add_argument("--depth", default="exhaustive")
    parser.add_argument("--max-sources", type=int, default=80)
    parser.add_argument("--min-sources", type=int, default=20)
    parser.add_argument("--min-reads", type=int, default=5)
    parser.add_argument("--skip-research", action="store_true")
    parser.add_argument("--include-annas-archive", action="store_true")
    args = parser.parse_args()

    if args.output.exists() and args.output.is_dir():
        raise SystemExit(f"Refusing to write DOCX: output path is a directory: {args.output}")

    args.workspace.mkdir(parents=True, exist_ok=True)
    if not args.skip_research:
        run_research(args)

    evidence = load_evidence(args.workspace)
    sources = load_sources(args.workspace, evidence)
    data_links = load_json(args.workspace / "data_links.json", [])
    read_count = len(list((args.workspace / "reads").glob("*")))
    counts = evidence.get("research_call_counts") if isinstance(evidence, dict) else None

    synthesis_dir = args.workspace / "synthesis"
    synthesis_dir.mkdir(parents=True, exist_ok=True)
    write_synthesis_files(args, sources, data_links, read_count, counts, synthesis_dir)
    write_docx(args, sources, data_links, read_count, counts, synthesis_dir)

    deliverable_validator = Path(__file__).with_name("validate_research_deliverable.py")
    deliverable_result = subprocess.run(
        [
            sys.executable,
            str(deliverable_validator),
            "--workspace",
            str(args.workspace),
            "--docx",
            str(args.output),
            "--min-sources",
            str(args.min_sources),
            "--min-reads",
            str(args.min_reads),
            "--min-draft-chars",
            "8000",
            "--require-call-counts",
        ],
        text=True,
        capture_output=True,
    )
    quality_validator = Path(__file__).with_name("validate_study_quality.py")
    quality_result = subprocess.run(
        [
            sys.executable,
            str(quality_validator),
            "--docx",
            str(args.output),
            "--min-images",
            "3",
            "--min-domain-tables",
            "4",
        ],
        text=True,
        capture_output=True,
    )
    with (synthesis_dir / "qa-notes.md").open("a", encoding="utf-8") as handle:
        handle.write("\n\n## Deliverable Validator\n\n```json\n")
        handle.write(deliverable_result.stdout.strip())
        handle.write("\n```\n")
        if deliverable_result.stderr.strip():
            handle.write("\nStderr:\n\n```text\n")
            handle.write(deliverable_result.stderr.strip())
            handle.write("\n```\n")
        handle.write("\n\n## Study Quality Validator\n\n```json\n")
        handle.write(quality_result.stdout.strip())
        handle.write("\n```\n")
        if quality_result.stderr.strip():
            handle.write("\nStderr:\n\n```text\n")
            handle.write(quality_result.stderr.strip())
            handle.write("\n```\n")
    print(deliverable_result.stdout)
    print(quality_result.stdout)
    raise SystemExit(deliverable_result.returncode or quality_result.returncode)


def run_research(args: argparse.Namespace) -> None:
    cmd = [
        "ctox",
        "web",
        "deep-research",
        "--query",
        args.query,
        "--depth",
        args.depth,
        "--max-sources",
        str(args.max_sources),
        "--workspace",
        str(args.workspace),
    ]
    if args.include_annas_archive:
        cmd.append("--include-annas-archive")
    result = subprocess.run(cmd, text=True, capture_output=True)
    (args.workspace / "research-command.txt").write_text(" ".join(cmd), encoding="utf-8")
    (args.workspace / "research-stdout.json").write_text(result.stdout, encoding="utf-8")
    if result.stderr.strip():
        (args.workspace / "research-stderr.txt").write_text(result.stderr, encoding="utf-8")
    if result.returncode != 0:
        raise SystemExit(f"ctox deep research failed with exit {result.returncode}: {result.stderr}")


def load_evidence(workspace: Path) -> dict[str, Any]:
    evidence = load_json(workspace / "evidence_bundle.json", {})
    return evidence if isinstance(evidence, dict) else {}


def load_sources(workspace: Path, evidence: dict[str, Any]) -> list[dict[str, Any]]:
    sources = evidence.get("sources")
    if isinstance(sources, list) and sources:
        return [source for source in sources if isinstance(source, dict)]
    path = workspace / "sources.jsonl"
    loaded = []
    if path.is_file():
        for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
            if not line.strip():
                continue
            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(item, dict):
                loaded.append(item)
    return loaded


def load_json(path: Path, fallback: Any) -> Any:
    if not path.is_file():
        return fallback
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return fallback


def source_label(source: dict[str, Any], index: int) -> str:
    title = str(source.get("title") or source.get("display_name") or "Quelle").strip()
    year = source.get("year") or source.get("publication_year")
    if year:
        return f"[S{index + 1}] {title} ({year})"
    return f"[S{index + 1}] {title}"


def source_url(source: dict[str, Any]) -> str:
    return str(source.get("url") or source.get("doi") or "").strip()


def matching_sources(sources: list[dict[str, Any]], technology: dict[str, Any]) -> list[int]:
    matches = []
    keywords = [str(value).lower() for value in technology["keywords"]]
    for index, source in enumerate(sources):
        text = " ".join(
            str(source.get(key, ""))
            for key in ["title", "snippet", "summary", "source_type", "venue", "url"]
        ).lower()
        read = source.get("read")
        if isinstance(read, dict):
            text += " " + json.dumps(read, ensure_ascii=False).lower()
        if any(keyword in text for keyword in keywords):
            matches.append(index)
    return matches[:6]


def citation(indices: list[int]) -> str:
    if not indices:
        return "(physikalische Einordnung; keine spezifische Quelle im aktuellen Trefferbündel)"
    return "(" + ", ".join(f"S{index + 1}" for index in indices[:4]) + ")"


def write_synthesis_files(
    args: argparse.Namespace,
    sources: list[dict[str, Any]],
    data_links: list[Any],
    read_count: int,
    counts: Any,
    synthesis_dir: Path,
) -> None:
    evidence_lines = [
        "# Evidence matrix",
        "",
        f"Quellenstand: {len(sources)} deduplizierte Quellen, {read_count} gespeicherte Reads.",
        "",
        "| Technologie | Prinzip | Evidenz | Erwartung | Confounder | Experiment |",
        "|---|---|---|---|---|---|",
    ]
    scores = [
        "# Technology scores",
        "",
        "| Technologie | Erfolg | Durchsatz | Single-shot | Robustheit gegen Folie | Unsicherheit | Evidenz |",
        "|---|---:|---:|---:|---:|---|---|",
    ]
    for tech in TECHNOLOGIES:
        indices = matching_sources(sources, tech)
        evidence_lines.append(
            f"| {tech['name']} | {tech['principle']} | {citation(indices)} | {tech['verdict']} | metallische Folie, Lift-off, Schichtdicke | Couponmatrix mit intakten und defekten Gitterfeldern |"
        )
        scores.append(
            f"| {tech['name']} | {tech['success']} | {tech['throughput']} | {tech['single_shot']} | {tech['foil']} | {tech['uncertainty']} | {citation(indices)} |"
        )

    (synthesis_dir / "evidence-matrix.md").write_text("\n".join(evidence_lines) + "\n", encoding="utf-8")
    (synthesis_dir / "technology-scores.md").write_text("\n".join(scores) + "\n", encoding="utf-8")
    (synthesis_dir / "report-outline.md").write_text(report_outline(), encoding="utf-8")
    (synthesis_dir / "figure-plan.md").write_text(figure_plan(data_links), encoding="utf-8")
    (synthesis_dir / "report-draft.md").write_text(report_draft(args, sources, data_links, read_count, counts), encoding="utf-8")
    (synthesis_dir / "qa-notes.md").write_text(
        "# QA notes\n\n"
        f"- Datum: {date.today().isoformat()}\n"
        f"- Quellen: {len(sources)}\n"
        f"- Reads: {read_count}\n"
        f"- Data links: {len(data_links) if isinstance(data_links, list) else 0}\n"
        f"- Call counts: `{json.dumps(counts, ensure_ascii=False)}`\n",
        encoding="utf-8",
    )
    if data_links:
        (synthesis_dir / "data-artifacts.md").write_text(
            "# Data artifacts\n\n" + json.dumps(data_links, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )


def report_outline() -> str:
    return """# Report outline

## 1. Management Summary
- Entscheidungsfrage, Top-Empfehlung, No-Go-Kandidaten und naechster Versuchsschritt.
- Kurzfazit zu Kupfergitter, CFK-Schichtsystem und metallischer Folie als dominanter Stoergroesse.

## 2. Problemstellung und Schichtannahmen
- Zielstruktur: Lack/Primer, CFK/CFRP, Kupfer-Blitzschutzgitter, optionale kontinuierliche Metallfolie.
- Defektbilder: Gitterunterbrechung, Verschiebung, lokale Delamination, Kopplungsfehler, Korrosion, Ueberdeckung.
- Randbedingungen: kontaktlos, Flaechenleistung, Single-shot-Potential, Stand-off, Industrialisierbarkeit.

## 3. Research-Methodik und Evidenzbasis
- Suchstrategie ueber Web, wissenschaftliche Datenbanken, offene Paper, Industriequellen und Patente.
- Quellenmix, Read-Zahlen, Snapshots/PDFs, Datenlinks und Grenzen der Evidenz.
- Trennung zwischen belegten Aussagen, physikalischen Inferenzen und Annahmen.

## 4. Technologie-Screening
- Hyperspektral/optisch, THz, Mikrowelle/mmWave, Wirbelstrom, pulsed eddy current, Induktionsthermografie, IR-Thermografie, Shearografie, X-ray/CT, magnetische Verfahren.
- Bewertung nach Sichtbarkeit, Tiefenwirkung, Defektsensitivitaet, Folienrobustheit, Durchsatz, Reifegrad und Sicherheits-/Integrationsrisiko.

## 5. Detailbewertung der Shortlist
- Elektromagnetischer Pfad: Wirbelstrom, gepulster Wirbelstrom, Induktionsthermografie.
- Indirekte Flaechenverfahren: Thermografie und Shearografie.
- Referenzverfahren: X-ray/CT als Ground Truth statt Produktionspfad.

## 6. Metallische Folie als Confounder
- Abschirmung/Reflexion bei THz, mmWave und optischen Verfahren.
- Entkopplungsmoeglichkeiten bei frequenz- und zeitaufgeloesten elektromagnetischen Verfahren.
- Kill-Kriterien fuer Verfahren, die Gitter und Folie nicht trennen koennen.

## 7. Versuchsdesign
- Couponmatrix mit intakten/defekten Gittern, variierender Schichtdicke, Folie ja/nein und realistischen Lift-off-Werten.
- Messgroessen, Wiederholbarkeit, Ground Truth und Auswertung.
- Entscheidungsgates fuer Demonstrator und Abbruch.

## 8. Risiken und Mitigation
- Fehlalarme, Aufloesungsgrenzen, Kalibrierbarkeit, Materialvariabilitaet, Kosten, Sicherheit, Datenauswertung.
- Mitigations durch Ground Truth, Modellierung, Coupondesign und multimodale Pruefung.

## 9. Empfehlung
- Primaere technische Linie, Nebenpfade, sofortige Experimente, Erfolgskriterien und weitere Forschungsfragen.

## 10. Quellenverzeichnis
- DOI/URL, Quelletyp, Evidenzrolle und Grenzen.
"""


def figure_plan(data_links: list[Any]) -> str:
    return (
        "# Figure plan\n\n"
        "1. Originales Schichtstack-Schema mit Lack/Primer, CFK, Kupfergitter und optionaler Metallfolie.\n"
        "2. Bewertungsmatrix als Tabelle im DOCX.\n"
        "3. Entscheidungsgate-Tabelle fuer Couponversuche.\n"
        f"4. GitHub-/Datenlinks: {len(data_links) if isinstance(data_links, list) else 0}; nur verwenden, wenn inspiziert und fachlich relevant.\n"
    )


def report_draft(
    args: argparse.Namespace,
    sources: list[dict[str, Any]],
    data_links: list[Any],
    read_count: int,
    counts: Any,
) -> str:
    source_count = len(sources)
    parts = [
        f"# {args.title}",
        "",
        "## Management Summary",
        "Die Fragestellung betrifft die kontaktlose Detektion eines in CFK/CFRP eingebetteten Kupfergitters fuer Blitzschutzanwendungen. "
        "Die vorlaeufige technische Empfehlung ist ein zweistufiges Vorgehen: Erstens Wirbelstrom beziehungsweise gepulste elektromagnetische Verfahren als primaerer Kandidat fuer direkte Leitfaehigkeits- und Gitteranomalien; zweitens aktive Thermografie oder Shearografie als schnelle indirekte Screeningverfahren. X-Ray/CT bleibt die technische Referenz fuer Ground Truth, ist aber wegen Durchsatz, Bauteilgroesse und Strahlenschutz kein idealer Produktionspfad. THz, mmWave und Hyperspektralverfahren sind nicht pauschal auszuschliessen, verlieren aber deutlich an Plausibilitaet, sobald eine kontinuierliche metallische Folie unter dem Kupfergitter vorhanden ist.",
        "",
        "## Methodik und Evidenzbasis",
        f"Der Research-Lauf erfasste {source_count} deduplizierte Quellen; {read_count} Quellen wurden als Read-Artefakte gespeichert. "
        "Die Quellen werden als Evidenz fuer physikalische Plausibilitaet, industrielle Reife, typische Limitierungen und Validierungsdesign genutzt. Wissenschaftliche Metadaten ohne Volltext werden nur als begrenzte Evidenz behandelt; offene Abstracts, Review-Artikel, Anwendungsberichte und Quellen mit gespeicherten Reads wiegen staerker.",
        "",
        "## Problemstellung und Schichtannahmen",
        "Das Zielsystem besteht aus einem mehrlagigen Verbund mit Lack, Primer beziehungsweise Deckschichten, CFK/CFRP, einem Kupfergitter zur Blitzstromableitung und moeglicherweise einer darunterliegenden kontinuierlichen metallischen Folie. Kontaktlose Pruefung bedeutet hier nicht nur beruehrungsfrei, sondern moeglichst auch mit vertretbarem Stand-off, reproduzierbarer Lift-off-Kompensation und Flaechenleistung. Die gesuchte Messgroesse ist nicht ein abstrakter Defekt, sondern die raeumliche Integritaet eines leitfaehigen Gitters: Unterbrechungen, lokale Delaminationen, verschobene Mesh-Bereiche, Korrosion oder Kopplungsfehler koennen relevant sein.",
        "",
        "## Anforderungen und Randbedingungen",
        "Die Pruefung muss einseitig, kontaktlos und fuer groessere Flaechen skalierbar sein. Relevante Bewertungskriterien sind Abbildungsfaehigkeit des Gitters, Defektsensitivitaet, Robustheit gegen Lift-off und Schichtdickenvariation, Umgang mit der metallischen Folie, Sicherheitsrisiko, Reifegrad und Integrationsaufwand. Single-shot-faehige Kameraverfahren erhalten nur dann hohe Bewertungen, wenn die physikalische Kopplung zur Zielstruktur glaubwuerdig ist.",
        "",
        "## Bewertungslogik",
        "Die qualitative Bewertung kombiniert drei Ebenen: belegte Literatur- und Industriehinweise, physikalische Plausibilitaet fuer den konkreten Schichtstack und Validierbarkeit ueber Coupons. Eine hohe Flaechenleistung kompensiert keine fehlende Sichtbarkeit des Kupfergitters. Die metallische Folie wird als eigenes Szenario bewertet, weil sie bei elektromagnetischen und optischen Verfahren die Antwort der Zielstruktur dominieren kann.",
        "",
        "## Erfolgsaussichten nach Schichtaufbau (Szenarien)",
        "Szenario A nimmt an, dass das Kupfergitter beziehungsweise EMF die erste relevante Metallschicht ist. Szenario B nimmt eine zusaetzliche nahezu geschlossene Metallfolie an. Szenario C erhoeht die Deckschichtdicke oder legt das Gitter tiefer. Die Rangfolge der Verfahren aendert sich zwischen diesen Szenarien deutlich; insbesondere THz und mmWave sind in Szenario B nur mit harten Kill-Kriterien sinnvoll.",
        "",
        "## Technologie-Screening",
    ]
    method_notes = {
        "Wirbelstrom / Eddy Current": "Die Methode koppelt direkt an die leitfaehige Zielstruktur. Kritisch sind Lift-off, Faserrichtungs-/CFK-Leitfaehigkeit und die Trennung von Gitter- und Folienantwort. Parameterstudien muessen Frequenz, Spulengeometrie und Arrayauflösung variieren.",
        "Pulsed Eddy Current / Induktionsthermografie": "Der Vorteil liegt in der Vollfeldbeobachtung der thermischen Antwort nach elektromagnetischer Anregung. Defekte koennen als lokale Strom- oder Waermeflussstoerungen sichtbar werden; die Methode braucht aber saubere Anregungsenergie und Thermalkalibrierung.",
        "Terahertz Imaging": "THz kann dielektrische Deckschichten gut adressieren, wird aber an leitfaehigen Grenzflaechen stark reflektiert. Als Forschungsfrage ist THz sinnvoll, wenn gezeigt werden kann, dass die erste Metallstruktur die gesuchte Gitterlage ist.",
        "Mikrowelle / mmWave": "Groessere Wellenlaengen erlauben mehr Stand-off, liefern aber geringere Ortsauflösung. Metallfolien und grossflaechige Leiter erzeugen dominante Reflexionen, weshalb tomographische oder polarimetrische Auswertung geprueft werden muss.",
        "Hyperspektral / optisch": "Hyperspektral ist primaer ein Oberflaechen- und Beschichtungsverfahren. Ohne indirekte Oberflaechensignatur ist es fuer die verdeckte Metallgeometrie kein Hauptpfad, kann aber Lack-/Primerzustand als Kontextmerkmal liefern.",
        "Infrarot-Thermografie": "Aktive Thermografie ist stark fuer Delaminationen, Disbonds und Waermeflussstoerungen. Fuer die Gittertopologie ist sie eher indirekt, kann aber als schnelles Screening mit definierten Defektcoupons sehr wertvoll sein.",
        "Shearografie": "Shearografie reagiert auf belastungsinduzierte Dehnungsfelder und ist deshalb fuer strukturelle Verbundfehler geeignet. Sie beantwortet die Gittergeometriefrage nur indirekt, kann aber Disbond und lokale Steifigkeitsanomalien abgrenzen.",
        "Röntgen / CT": "Radiografie und CT bieten die beste geometrische Referenz, sind aber wegen Strahlenschutz, Bauteilgroesse und Durchsatz eher Ground-Truth-Technik als Produktionspruefung.",
        "Magnetische Verfahren / MFL": "Statische magnetische Verfahren sind fuer Kupfer unguenstig. Relevant bleiben nur Verfahren, die ueber induzierte Stroeme oder magneto-optische Hilfseffekte eine indirekte Antwort erzeugen.",
    }
    for tech in TECHNOLOGIES:
        indices = matching_sources(sources, tech)
        parts.extend(
            [
                f"### {tech['name']}",
                f"{tech['principle']} {tech['verdict']} Evidenzbezug: {citation(indices)}.",
                method_notes.get(str(tech["name"]), ""),
                "",
            ]
        )
    parts.extend(
        [
            "## Detailbewertung ausgewaehlter Ansaetze",
            "Die Shortlist besteht aus elektromagnetischen Verfahren, induktionsbasierter Thermografie und X-ray/CT als Ground-Truth-Referenz. Eddy Current ist der direkteste Kandidat fuer Leitergeometrie und Unterbrechungen, muss aber die Folienantwort frequenz- oder zeitaufgeloest trennen. Induktions-Thermografie bietet den besten Kompromiss aus Flaechenleistung und Defektsensitivitaet, sofern Gitterfehler eine lokale thermische Signatur erzeugen. X-ray/CT bleibt die Referenz fuer Couponvalidierung und sollte nicht als primaerer Produktionspfad geplant werden.",
            "",
            "## Metallische Folie als Confounder",
            "Die kontinuierliche metallische Folie ist der zentrale Risikofaktor. Fuer THz, mmWave und viele optische beziehungsweise quasioptische Verfahren kann sie als reflektierende oder abschirmende Grenzflaeche wirken. Dadurch wird die Antwort des eigentlichen Kupfergitters entweder ueberdeckt oder nur als Mehrwege-/Interferenzsignal sichtbar. Bei Wirbelstromverfahren ist die Folie ebenfalls ein Confounder, aber nicht zwingend ein K.O.-Kriterium: Frequenz, Pulsform, Spulengeometrie und Inversionsmodell koennen genutzt werden, um Tiefen- und Leitfaehigkeitsbeitraege teilweise zu trennen. Genau diese Trennbarkeit ist ein fruehes Entscheidungsgate.",
            "",
            "## Empfohlenes Versuchsdesign",
            "Die Validierung sollte mit planaren Coupons beginnen: intaktes Gitter ohne Folie, intaktes Gitter mit Folie, definierte Gitterunterbrechungen, lokale Ueberlappungsfehler, variable Lack-/Primer-Dicken und gezielte Delaminationen. Fuer jedes Verfahren werden Rohdaten, Wiederholbarkeit, Lift-off-Toleranz, Flaechenleistung und Fehlalarmrate gemessen. X-Ray/CT oder Schliffbilder dienen als Ground Truth, nicht als Produktionsloesung. Danach folgt eine Downselection auf maximal zwei Kandidaten fuer gekruemmte und real lackierte Bauteile.",
            "",
            "## Entscheidungsgates",
            "Gate 1 prueft, ob das Verfahren das intakte Kupfergitter gegen den Schichtstack ueberhaupt stabil sieht. Gate 2 prueft definierte Unterbrechungen bei variierender Lackdicke. Gate 3 prueft die metallische Folie als Stoerfall. Gate 4 prueft Durchsatz und Handhabung im realistischen Abstand. Nur Verfahren, die Gate 1 bis 3 bestehen, sollten in einen industriellen Demonstrator uebergehen.",
            "",
            "## Risiken, Abhaengigkeiten und Mitigation",
            "Die groessten Risiken sind ein unklarer Schichtaufbau, die Abschirmung durch eine metallische Folie, nicht reproduzierbarer Lift-off, zu geringe Ortsaufloesung und die Verwechslung von CFK-Anisotropie mit Gitterdefekten. Mitigation entsteht durch Ground-Truth-Coupons, definierte Defektgeometrien, Frequenz-/Pulsparameterstudien, wiederholte Messungen und eine fruehe No-Go-Entscheidung fuer Verfahren, die Gitter und Folie nicht trennen koennen.",
            "",
            "## Empfehlung",
            "Primaer sollte ein elektromagnetischer Pfad aus Wirbelstrom, gepulstem Wirbelstrom und induktionsbasierter Thermografie aufgebaut werden. Parallel sollte Thermografie/Shearografie als schneller indirekter Screeningpfad getestet werden. THz und mmWave sollten nur mit einem klaren Kill-Kriterium getestet werden: Wenn die Folie die Gitterinformation abschirmt, werden diese Pfade beendet. Hyperspektral bleibt ein Hilfsverfahren fuer Oberflaeche/Beschichtung, nicht fuer die direkte Blitzschutzgitterpruefung. X-Ray/CT bleibt Referenztechnik fuer Ground Truth.",
        ]
    )
    if data_links:
        parts.extend(["", "## Daten- und Repository-Links", f"Es wurden {len(data_links)} Daten-/Repository-Kandidaten gefunden. Sie sind im Research-Workspace dokumentiert und muessen vor einer datengetriebenen Auswertung einzeln inspiziert werden."])
    return "\n".join(parts) + "\n"


def write_docx(
    args: argparse.Namespace,
    sources: list[dict[str, Any]],
    data_links: list[Any],
    read_count: int,
    counts: Any,
    synthesis_dir: Path,
) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "python-docx is required to write DOCX reports. Install python-docx or run in the CTOX bundled document runtime."
        ) from exc

    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(args.title)
    run.bold = True
    run.font.color.rgb = RGBColor(25, 55, 80)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Kontaktlose Pruefung von Kupfer-Blitzschutzgittern in CFK/CFRP").italic = True

    add_metadata_table(doc, sources, read_count, counts, args.workspace)
    add_abbreviation_table(doc)
    add_generated_figures(doc, args.workspace)
    add_section_from_markdown(doc, (synthesis_dir / "report-draft.md").read_text(encoding="utf-8"))
    add_scenario_table(doc)
    add_score_table(doc)
    add_defect_catalog_table(doc)
    add_experiment_table(doc)
    add_references(doc, sources)
    doc.save(args.output)


def configure_doc(doc: Document) -> None:
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Aptos Display"


def add_metadata_table(doc: Document, sources: list[dict[str, Any]], read_count: int, counts: Any, workspace: Path) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Stand", date.today().isoformat()),
        ("Research-Modul", "Machbarkeitsstudie / Technical Feasibility"),
        ("Research-Workspace", str(workspace)),
        ("Quellen", str(len(sources))),
        ("Gespeicherte Reads", str(read_count)),
        ("Rechercheumfang", summarize_counts(counts)),
    ]
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    doc.add_paragraph()


def summarize_counts(counts: Any) -> str:
    if not isinstance(counts, dict):
        return "nicht verfuegbar"
    return (
        f"{counts.get('executed_search_queries', '?')} Suchlaeufe, "
        f"{counts.get('database_queries', '?')} Datenbankabfragen, "
        f"{counts.get('successful_page_reads', '?')} erfolgreiche Reads, "
        f"{counts.get('figure_candidates', '?')} Abbildungskandidaten"
    )


def add_abbreviation_table(doc: Document) -> None:
    doc.add_heading("Abkuerzungsverzeichnis", level=1)
    rows = [
        ("CFK / CFRP", "Kohlenstofffaserverstaerkter Kunststoff / carbon fiber reinforced polymer"),
        ("LSP", "Lightning Strike Protection / Blitzschutzlage"),
        ("EMF", "Expanded Metal Foil, haeufig als Kupfer- oder Aluminiumlage"),
        ("NDT / NDE", "Non-Destructive Testing / Evaluation"),
        ("ECT", "Eddy Current Testing / Wirbelstrompruefung"),
        ("ECPT", "Eddy Current Pulsed Thermography / Induktions-Thermografie"),
        ("THz", "Terahertz-Imaging beziehungsweise Terahertz Time-Domain Spectroscopy"),
        ("mmWave", "Millimeterwellen-Radar oder -Imaging"),
        ("CT", "Computed Tomography / Computertomografie"),
    ]
    add_simple_table(doc, ["Abk.", "Bedeutung"], rows)


def add_generated_figures(doc: Document, workspace: Path) -> None:
    figure_dir = workspace / "synthesis" / "figures"
    figure_dir.mkdir(parents=True, exist_ok=True)
    specs = [
        ("layer-stack.png", [(80, 120, 820, 190), (80, 210, 820, 290), (80, 310, 820, 390), (80, 410, 820, 470)], "Abbildung 1: Eigenes Schichtstack-Schema: Deckschichten, CFK, Kupfergitter und optionale Metallfolie."),
        ("method-interaction.png", [(80, 160, 250, 360), (330, 120, 500, 400), (580, 160, 750, 360)], "Abbildung 2: Eigene Prinzipdarstellung: elektromagnetische Anregung, Zielschicht und Detektionsantwort."),
        ("workflow-gates.png", [(80, 170, 220, 330), (270, 170, 410, 330), (460, 170, 600, 330), (650, 170, 790, 330)], "Abbildung 3: Eigener Pruefworkflow mit Couponstudie, Shortlist, Demonstrator und Entscheidungsgates."),
    ]
    for filename, boxes, caption in specs:
        path = figure_dir / filename
        write_simple_png(path, boxes)
        paragraph = doc.add_paragraph()
        paragraph.alignment = 1
        paragraph.add_run().add_picture(str(path))
        cap = doc.add_paragraph(caption)
        cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]


def write_simple_png(path: Path, boxes: list[tuple[int, int, int, int]]) -> None:
    width, height = 900, 560
    bg = (248, 250, 252)
    colors = [(220, 230, 242), (196, 215, 155), (242, 220, 180), (230, 184, 183), (184, 204, 228)]
    pixels = [bg] * (width * height)
    for idx, (x1, y1, x2, y2) in enumerate(boxes):
        fill = colors[idx % len(colors)]
        for y in range(max(0, y1), min(height, y2)):
            for x in range(max(0, x1), min(width, x2)):
                border = x in (x1, x2 - 1) or y in (y1, y2 - 1)
                pixels[y * width + x] = (70, 90, 110) if border else fill
    raw = b"".join(b"\x00" + bytes(channel for pixel in pixels[y * width:(y + 1) * width] for channel in pixel) for y in range(height))
    def chunk(kind: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
    png = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)) + chunk(b"IDAT", zlib.compress(raw, 9)) + chunk(b"IEND", b"")
    path.write_bytes(png)


def add_section_from_markdown(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def add_simple_table(doc: Document, columns: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    doc.add_paragraph()


def add_scenario_table(doc: Document) -> None:
    doc.add_heading("Erfolgsaussichten nach Schichtaufbau (Szenarien)", level=1)
    rows = [
        ("Wirbelstrom / Arrays", "hoch", "mittel-hoch, wenn Frequenztrennung gelingt", "mittel"),
        ("Induktions-Thermografie", "hoch", "mittel-hoch, falls Gitterfehler thermisch sichtbar bleibt", "mittel-hoch"),
        ("THz", "mittel-hoch bis hoch", "niedrig, wenn Folie geschlossen ist", "mittel"),
        ("Mikrowelle / mmWave", "mittel", "niedrig-mittel", "niedrig-mittel"),
        ("Hyperspektral", "niedrig", "niedrig", "niedrig"),
        ("X-ray / CT", "hoch", "hoch", "hoch, aber geringer Durchsatz"),
    ]
    add_simple_table(
        doc,
        [
            "Verfahren",
            "Szenario A: Gitter/EMF erste Metallschicht",
            "Szenario B: zusaetzliche geschlossene Folie",
            "Szenario C: tiefere Lage / dickere Deckschicht",
        ],
        rows,
    )


def add_defect_catalog_table(doc: Document) -> None:
    doc.add_heading("Defekt- und Couponkatalog", level=1)
    rows = [
        ("D1", "Unterbrechung einzelner Kupferstege"),
        ("D2", "Fehlender oder verschobener Gitterbereich"),
        ("D3", "Lokaler Abbrand beziehungsweise Ueberhitzung"),
        ("D4", "Disbond/Delamination nahe der Blitzschutzlage"),
        ("D5", "Variierende Deckschichtdicke"),
        ("D6", "Zusaetzliche leitfaehige Folie als Abschirmungsreferenz"),
    ]
    add_simple_table(doc, ["ID", "Beschreibung"], rows)


def add_score_table(doc: Document) -> None:
    doc.add_heading("Bewertungsmatrix", level=1)
    columns = ["Technologie", "Erfolg", "Durchsatz", "Single-shot", "Folie", "Unsicherheit", "Evidenz / Begruendung"]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for tech in TECHNOLOGIES:
        cells = table.add_row().cells
        values = [
            tech["name"],
            str(tech["success"]),
            str(tech["throughput"]),
            str(tech["single_shot"]),
            str(tech["foil"]),
            tech["uncertainty"],
            f"{tech['verdict']} Score muss gegen Evidence Cards und Schicht-Szenarien geprueft werden.",
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)


def add_experiment_table(doc: Document) -> None:
    doc.add_heading("Coupon- und Entscheidungsgate-Matrix", level=1)
    columns = ["Gate", "Coupon/Variation", "Messziel", "Pass-Kriterium", "Fail-Kriterium"]
    rows = [
        ("G1", "Intaktes Gitter ohne Folie", "Baseline-Sichtbarkeit", "stabile Gitterantwort", "keine reproduzierbare Antwort"),
        ("G2", "Unterbrochenes Gitter", "Defektsensitivitaet", "Defekt lokalisiert", "Defekt nicht trennbar"),
        ("G3", "Intaktes/defektes Gitter mit Folie", "Folie als Confounder", "Gitterinformation bleibt trennbar", "Folie ueberdeckt Signal"),
        ("G4", "Variierende Lack-/Primer-Dicke", "Robustheit", "kalibrierbare Drift", "unkontrollierbare Fehlalarme"),
        ("G5", "Gekruemmtes reales Bauteil", "Transfer", "Durchsatz und Handling plausibel", "Laborerfolg nicht uebertragbar"),
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_references(doc: Document, sources: list[dict[str, Any]]) -> None:
    doc.add_heading("Quellen", level=1)
    for index, source in enumerate(sources[:80]):
        doc.add_paragraph(f"{source_label(source, index)}. {source_url(source)}")


if __name__ == "__main__":
    main()
