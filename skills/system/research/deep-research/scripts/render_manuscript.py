#!/usr/bin/env python3
"""
render_manuscript.py — turn a deep-research manuscript JSON (on stdin)
into a DOCX file using python-docx.

Invoked by the Rust render module as a subprocess; can also be run by
the operator directly for debugging:

    ctox report render RUN_ID --format json --out manuscript.json
    cat manuscript.json | python3 render_manuscript.py --out report.docx

The manuscript JSON shape is documented in
`references/check_contracts.md` and `references/setup_guide.md`. Top
keys: manifest, title, subtitle, version_line, context_line,
scope_disclaimer, abbreviations, docs[], references[], figures[].

The script:
  - Loads python-docx (emits an explicit install hint and exit 2 if
    missing).
  - Sets A4 page, 2.5 cm margins, Arial 11 default.
  - Overrides Heading 1/2/3 styles for consistent black sans-serif
    headings (Word and LibreOffice render the same).
  - Renders title block, an inserted TOC field, and per-doc / per-block
    content following the kinds described in the contract:
    narrative, matrix, scenario_grid, risk_register, evidence_register,
    abbreviation_table, defect_catalog, competitor_matrix,
    criteria_table.
  - Lints the produced text against the asset_pack's
    forbidden_meta_phrases[] (warning to stderr, never blocks).
  - Replaces Unicode hyphens / dashes with ASCII equivalents
    (style_guidance: ASCII hyphens only).
  - Prints `OK <byte_count> <output_path>` on success and exits 0.

No global pip installs. No network. Runs on macOS dev machines and on
Ubuntu Linux production hosts (Yoda).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

ASSET_PACK_RELATIVE = Path(__file__).resolve().parent.parent / "references" / "asset_pack.json"

# Unicode hyphen / dash characters to normalize to ASCII.
# style_guidance: ASCII hyphens only.
_UNICODE_DASHES = {
    "‐": "-",  # hyphen
    "‑": "-",  # non-breaking hyphen
    "‒": "-",  # figure dash
    "–": "-",  # en dash
    "—": "-",  # em dash
    "―": "-",  # horizontal bar
    "′": "'",  # prime
    "″": '"',  # double prime
}


def _missing_python_docx() -> None:
    sys.stderr.write(
        "python-docx not installed. Run: python3 -m pip install python-docx\n"
    )
    sys.exit(2)


def _load_python_docx():
    """Import python-docx lazily so the missing-dep path is clean."""
    try:
        import docx  # type: ignore
        from docx.shared import Pt, Cm, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.enum.table import WD_ALIGN_VERTICAL  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
    except ImportError:
        _missing_python_docx()
        raise  # unreachable
    return {
        "docx": docx,
        "Pt": Pt,
        "Cm": Cm,
        "RGBColor": RGBColor,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_ALIGN_VERTICAL": WD_ALIGN_VERTICAL,
        "qn": qn,
        "OxmlElement": OxmlElement,
    }


def _ascii_dashes(text: str) -> str:
    if not text:
        return ""
    out = []
    for ch in text:
        out.append(_UNICODE_DASHES.get(ch, ch))
    return "".join(out)


def _load_forbidden_phrases() -> list:
    """Best-effort load of asset_pack.style_guidance.forbidden_meta_phrases.

    Missing or unparseable asset pack -> empty list (lint pass becomes a
    no-op). This keeps the renderer robust on hosts where the asset pack
    is at a different relative path.
    """
    try:
        with ASSET_PACK_RELATIVE.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
    except (OSError, json.JSONDecodeError):
        return []
    sg = data.get("style_guidance", {})
    phrases = sg.get("forbidden_meta_phrases", []) or []
    return [p for p in phrases if isinstance(p, str) and p.strip()]


def _override_heading_styles(document, deps) -> None:
    """Override Heading1/2/3 styles to Arial, black, with a consistent
    size cadence: 16/14/12 pt. Word and LibreOffice both honour the
    style.id values 'Heading1' / 'Heading2' / 'Heading3'.
    """
    Pt = deps["Pt"]
    RGBColor = deps["RGBColor"]
    cadence = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 12}
    for name, size in cadence.items():
        try:
            style = document.styles[name]
        except KeyError:
            continue
        font = style.font
        font.name = "Arial"
        font.size = Pt(size)
        font.bold = True
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _set_default_font(document, deps) -> None:
    """Set Normal style to Arial 11."""
    Pt = deps["Pt"]
    try:
        normal = document.styles["Normal"]
    except KeyError:
        return
    normal.font.name = "Arial"
    normal.font.size = Pt(11)


def _set_page(document, deps) -> None:
    """A4 with 2.5 cm margins on every section."""
    Cm = deps["Cm"]
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _add_inline_runs(paragraph, text: str, references: list, deps) -> None:
    """Render markdown-ish inline syntax into runs:
       **bold**, _italic_, [evidence-marker] -> superscript [N].

    `references` is the manuscript-level references list; an
    `evidence-marker` like `[e_017]` or `[REF-3]` looked up case-
    insensitively against `ref.evidence_id` / `ref.id` / `ref.ref_n`.
    Unknown markers render as the literal text.
    """
    Pt = deps["Pt"]
    text = _ascii_dashes(text or "")
    if not text:
        paragraph.add_run("")
        return

    # tokenize: ** ... **, _ ... _, [ ... ]
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|_[^_]+_|\[[A-Za-z0-9_\-:]+\])"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("[") and token.endswith("]"):
            marker = token[1:-1]
            ref_n = _lookup_ref_n(marker, references)
            if ref_n is not None:
                run = paragraph.add_run(f"[{ref_n}]")
                run.font.superscript = True
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _lookup_ref_n(marker: str, references: list):
    if not references:
        return None
    m = marker.strip().lower()
    for ref in references:
        for key in ("evidence_id", "id"):
            v = ref.get(key)
            if isinstance(v, str) and v.lower() == m:
                return ref.get("ref_n") or ref.get("number") or _index_of(ref, references)
        ref_n = ref.get("ref_n") or ref.get("number")
        if ref_n is not None and str(ref_n) == m:
            return ref_n
    return None


def _index_of(ref, references) -> int:
    for idx, r in enumerate(references, start=1):
        if r is ref:
            return idx
    return 0


def _add_paragraph_with_runs(document, text: str, references: list, deps,
                             style: str = None, alignment=None):
    p = document.add_paragraph(style=style) if style else document.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    _add_inline_runs(p, text, references, deps)
    return p


def _render_narrative(document, block: dict, references: list, deps) -> None:
    """Render a `narrative` block: heading at block.level, then markdown
    body split into paragraphs, with bullet detection."""
    body = block.get("markdown") or ""
    body = _ascii_dashes(body)
    lines = body.split("\n")
    bullet_run = []

    def flush_bullets():
        if not bullet_run:
            return
        for item in bullet_run:
            p = document.add_paragraph(style="List Bullet")
            _add_inline_runs(p, item, references, deps)
        bullet_run.clear()

    paragraph_lines: list = []

    def flush_para():
        if not paragraph_lines:
            return
        joined = " ".join(paragraph_lines).strip()
        if joined:
            p = document.add_paragraph()
            _add_inline_runs(p, joined, references, deps)
        paragraph_lines.clear()

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_para()
            flush_bullets()
            continue
        m_bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if m_bullet:
            flush_para()
            bullet_run.append(m_bullet.group(1))
            continue
        m_ord = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_ord:
            flush_para()
            flush_bullets()
            p = document.add_paragraph(style="List Number")
            _add_inline_runs(p, m_ord.group(2), references, deps)
            continue
        flush_bullets()
        paragraph_lines.append(stripped)
    flush_para()
    flush_bullets()


def _render_table_block(document, block: dict, references: list, deps) -> None:
    """Render a generic kind-with-table block: matrix, competitor_matrix,
    criteria_table, abbreviation_table, defect_catalog, scenario_grid,
    risk_register, evidence_register. Layout is a real Word table with a
    bold header row.
    """
    Pt = deps["Pt"]
    WD_ALIGN_VERTICAL = deps["WD_ALIGN_VERTICAL"]
    table_data = block.get("table") or {}
    headers = table_data.get("headers") or []
    rows = table_data.get("rows") or []
    if not headers:
        # Fall back to narrative if table data is missing.
        return _render_narrative(document, block, references, deps)
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for idx, hdr in enumerate(headers):
        cell = hdr_cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(_ascii_dashes(str(hdr)))
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row[: len(headers)]):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, str(value), references, deps)
    document.add_paragraph()  # spacer


def _render_evidence_register(document, references: list, deps) -> None:
    """Render the Anhang/evidence register as a numbered list. Each
    entry: ref_n, authors (year). title. venue. URL/DOI as hyperlink.
    """
    if not references:
        return
    for idx, ref in enumerate(references, start=1):
        ref_n = ref.get("ref_n") or ref.get("number") or idx
        authors = ref.get("authors") or ""
        if isinstance(authors, list):
            authors = "; ".join(str(a) for a in authors if a)
        year = ref.get("year")
        title = ref.get("title") or ""
        venue = ref.get("venue") or ""
        url = ref.get("url") or ref.get("source_url") or ""
        doi = ref.get("doi") or ""
        bits = [f"[{ref_n}]"]
        if authors:
            tail = f"{authors}"
            if year:
                tail += f" ({year})"
            tail += "."
            bits.append(tail)
        if title:
            bits.append(f"{title}.")
        if venue:
            bits.append(f"{venue}.")
        if doi:
            bits.append(f"DOI: {doi}")
        elif url:
            bits.append(url)
        text = " ".join(bits)
        p = document.add_paragraph()
        _add_inline_runs(p, _ascii_dashes(text), [], deps)


def _add_toc_field(document, deps) -> None:
    """Insert a Word TOC field. Word will populate on first open;
    LibreOffice fills it on conversion. Add a hint paragraph.
    """
    qn = deps["qn"]
    OxmlElement = deps["OxmlElement"]
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Inhaltsverzeichnis (mit Rechtsklick aktualisieren)."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)
    hint = document.add_paragraph()
    hint_run = hint.add_run(
        "Hinweis: Falls das Inhaltsverzeichnis leer erscheint, "
        'mit Rechtsklick -> "Feld aktualisieren" aktualisieren.'
    )
    hint_run.italic = True


def _render_block(document, block: dict, references: list, deps) -> None:
    title = block.get("title") or ""
    level = block.get("level")
    if level is None:
        level = 2
    try:
        level_int = int(level)
    except (TypeError, ValueError):
        level_int = 2
    level_int = max(1, min(level_int, 9))
    document.add_heading(_ascii_dashes(str(title)), level=level_int)
    kind = block.get("kind") or "narrative"
    if kind == "narrative":
        _render_narrative(document, block, references, deps)
    elif kind in (
        "matrix",
        "competitor_matrix",
        "criteria_table",
        "abbreviation_table",
        "defect_catalog",
        "scenario_grid",
        "risk_register",
    ):
        # If the block carries narrative markdown alongside the table,
        # render the markdown first as intro prose.
        if block.get("markdown"):
            _render_narrative(document, block, references, deps)
        _render_table_block(document, block, references, deps)
    elif kind == "evidence_register":
        # Inline evidence_register: render references inline rather
        # than waiting for the global Anhang.
        _render_evidence_register(document, references, deps)
    else:
        # Unknown kind -> treat as narrative, do not crash.
        _render_narrative(document, block, references, deps)


def _render_title_block(document, manuscript: dict, deps) -> None:
    Pt = deps["Pt"]
    WD_ALIGN_PARAGRAPH = deps["WD_ALIGN_PARAGRAPH"]
    title = _ascii_dashes(manuscript.get("title") or "")
    subtitle = _ascii_dashes(manuscript.get("subtitle") or "")
    version_line = _ascii_dashes(manuscript.get("version_line") or "")
    context_line = _ascii_dashes(manuscript.get("context_line") or "")
    scope_disclaimer = _ascii_dashes(manuscript.get("scope_disclaimer") or "")

    if title:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(20)
    if subtitle:
        p = document.add_paragraph()
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(13)
    if version_line:
        p = document.add_paragraph()
        run = p.add_run(version_line)
        run.font.size = Pt(10)
    if context_line:
        p = document.add_paragraph()
        run = p.add_run(context_line)
        run.font.size = Pt(10)
    if scope_disclaimer:
        document.add_paragraph()
        p = document.add_paragraph()
        run = p.add_run(scope_disclaimer)
        run.font.size = Pt(10)
        run.italic = True


def _render_abbreviations(document, manuscript: dict, references: list, deps) -> None:
    abbreviations = manuscript.get("abbreviations") or []
    if not abbreviations:
        return
    document.add_heading("Abkuerzungsverzeichnis", level=1)
    table = document.add_table(rows=1 + len(abbreviations), cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    Pt = deps["Pt"]
    for col, label in enumerate(["Abkuerzung", "Bedeutung"]):
        hdr[col].text = ""
        p = hdr[col].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, entry in enumerate(abbreviations, start=1):
        cells = table.rows[r_idx].cells
        cells[0].text = _ascii_dashes(str(entry.get("abk") or entry.get("abbr") or ""))
        cells[1].text = _ascii_dashes(str(entry.get("meaning") or entry.get("definition") or ""))
    document.add_paragraph()


def _walk_text_for_phrases(document) -> str:
    """Collect every paragraph and table-cell text into one string for
    the forbidden-phrase lint pass. Reads only what we just wrote.
    """
    parts = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def _lint_forbidden_phrases(document, phrases: list) -> None:
    """Warn to stderr if any forbidden_meta_phrase appears. Never blocks
    the save.
    """
    if not phrases:
        return
    haystack = _walk_text_for_phrases(document).lower()
    hits = []
    for phrase in phrases:
        needle = phrase.lower()
        if needle and needle in haystack:
            hits.append(phrase)
    if hits:
        sys.stderr.write(
            "render_manuscript.py: WARNING — manuscript contains "
            f"{len(hits)} forbidden_meta_phrase hit(s): {hits!r}\n"
        )


def _render_doc(document, doc: dict, references: list, deps) -> None:
    blocks = doc.get("blocks") or []
    if not blocks:
        return
    doc_title = doc.get("title")
    if doc_title:
        document.add_heading(_ascii_dashes(str(doc_title)), level=1)
    blocks_sorted = sorted(blocks, key=lambda b: (b.get("ord", 0), b.get("instance_id") or ""))
    for block in blocks_sorted:
        _render_block(document, block, references, deps)


def render(manuscript: dict, output_path: Path, report_type: str, language: str) -> int:
    deps = _load_python_docx()
    docx_pkg = deps["docx"]
    document = docx_pkg.Document()

    _set_page(document, deps)
    _set_default_font(document, deps)
    _override_heading_styles(document, deps)

    _render_title_block(document, manuscript, deps)
    document.add_page_break()

    document.add_heading("Inhaltsverzeichnis", level=1)
    _add_toc_field(document, deps)
    document.add_page_break()

    _render_abbreviations(document, manuscript, manuscript.get("references") or [], deps)

    references = manuscript.get("references") or []
    docs = manuscript.get("docs") or []
    if not docs:
        # Tolerate flat manuscripts without a docs[] wrapper by treating
        # the top-level `blocks` field as a single doc.
        flat_blocks = manuscript.get("blocks") or []
        if flat_blocks:
            _render_doc(
                document,
                {"doc_id": "doc_main", "title": manuscript.get("title") or "", "blocks": flat_blocks},
                references,
                deps,
            )
    else:
        for doc in docs:
            _render_doc(document, doc, references, deps)

    if references:
        document.add_heading("Anhang — Quellen", level=1)
        _render_evidence_register(document, references, deps)

    forbidden = _load_forbidden_phrases()
    _lint_forbidden_phrases(document, forbidden)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path.stat().st_size


def _read_manuscript() -> dict:
    raw = sys.stdin.read()
    if not raw.strip():
        sys.stderr.write("render_manuscript.py: empty stdin\n")
        sys.exit(3)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"render_manuscript.py: bad JSON on stdin: {exc}\n")
        sys.exit(3)


def main(argv: list) -> int:
    parser = argparse.ArgumentParser(
        prog="render_manuscript.py",
        description="Render a deep-research manuscript JSON (stdin) to DOCX.",
    )
    parser.add_argument("--out", required=True, help="Path to the output DOCX file")
    parser.add_argument("--report-type", default=None, help="Optional report_type_id hint")
    parser.add_argument("--language", default="de", help="BCP-47 / short language tag")
    args = parser.parse_args(argv)

    manuscript = _read_manuscript()
    output_path = Path(args.out).expanduser().resolve()
    report_type = args.report_type or (
        manuscript.get("manifest", {}).get("report_type_id") or "feasibility_study"
    )
    byte_count = render(manuscript, output_path, report_type, args.language)
    sys.stdout.write(f"OK {byte_count} {output_path}\n")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
