#!/usr/bin/env python3
"""Quality gate for technical feasibility-study DOCX outputs."""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--min-images", type=int, default=3)
    parser.add_argument("--min-domain-tables", type=int, default=4)
    args = parser.parse_args()

    failures: list[str] = []
    warnings: list[str] = []

    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError:
        print(json.dumps({"ok": False, "failures": ["python-docx is required"]}, indent=2))
        raise SystemExit(2)

    if not args.docx.is_file():
        failures.append(f"DOCX missing: {args.docx}")
        emit(args, failures, warnings)

    doc = Document(str(args.docx))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    body = "\n".join(paragraphs)
    lower = body.lower()

    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip() and p.style and ("Heading" in p.style.name or "Title" in p.style.name)
    ]
    heading_blob = "\n".join(headings).lower()

    required_heading_terms = [
        ("management summary", ["management summary"]),
        ("problem framing", ["ausgangslage", "problemstellung", "fragestellung"]),
        ("component stack", ["bauteilaufbau", "schicht", "system / material"]),
        ("requirements", ["anforderungen", "randbedingungen", "boundary"]),
        ("technology screening", ["technologie-screening", "technology screening"]),
        ("scoring logic", ["bewertungslogik", "scoring logic"]),
        ("scenario matrix", ["szenario", "szenarien", "scenario"]),
        ("shortlist detail", ["detailbewertung", "shortlist"]),
        ("experimental design", ["versuchsdesign", "experimental design"]),
        ("risks", ["risiken", "risks"]),
        ("conclusion", ["fazit", "empfehlung", "recommendation"]),
    ]
    for label, terms in required_heading_terms:
        if not any(term in heading_blob for term in terms):
            failures.append(f"missing required study section: {label}")

    table_roles = classify_tables(doc)
    for role in ["abbreviations", "technology_matrix", "scenario_matrix", "defect_catalog"]:
        if role not in table_roles:
            failures.append(f"missing required domain table: {role}")
    if len(table_roles) < args.min_domain_tables:
        failures.append(
            f"only {len(table_roles)} domain table type(s) detected; expected at least {args.min_domain_tables}"
        )

    image_count = len(doc.inline_shapes)
    if image_count < args.min_images:
        failures.append(f"only {image_count} image/diagram(s); expected at least {args.min_images}")

    if "call counts:" in lower or "executed_search_queries" in lower or "estimated_external_fetches" in lower:
        failures.append("raw tool/call-count JSON appears in the narrative")
    if "evidenz noch schwach" in lower:
        failures.append("final report contains unresolved 'Evidenz noch schwach' markers")
    if "```" in body or re.search(r"^\s*[\{\[]\s*$", body, re.MULTILINE):
        failures.append("raw JSON/markdown code artefacts appear in the report body")

    duplicates = repeated_long_paragraphs(paragraphs)
    if duplicates:
        failures.append(
            "repeated generic long paragraph(s): "
            + "; ".join(text[:120] for text, _count in duplicates[:3])
        )

    source_dump_index = first_index(paragraphs, ["Quellenbasis", "Sources", "Quellen"])
    if source_dump_index is not None and source_dump_index < max(25, len(paragraphs) // 2):
        warnings.append("source/reference section appears early; verify report is not a source dump")

    if not score_rows_have_rationale(doc, heading_blob):
        failures.append("technology scoring lacks evidence/rationale support")

    emit(args, failures, warnings, doc, table_roles)


def classify_tables(doc) -> set[str]:
    roles: set[str] = set()
    for table in doc.tables:
        if not table.rows:
            continue
        headers = " | ".join(cell.text.strip().lower() for cell in table.rows[0].cells)
        if "abk" in headers and "bedeutung" in headers:
            roles.add("abbreviations")
        if ("verfahren" in headers or "technologie" in headers) and (
            "gitter" in headers or "erfolg" in headers or "fläche" in headers or "flaeche" in headers
        ):
            roles.add("technology_matrix")
        if "szenario" in headers or "scenario" in headers:
            roles.add("scenario_matrix")
        if ("id" in headers and "beschreibung" in headers) or "defekt" in headers or "coupon" in headers:
            roles.add("defect_catalog")
    return roles


def repeated_long_paragraphs(paragraphs: list[str]) -> list[tuple[str, int]]:
    normalized: list[str] = []
    original: dict[str, str] = {}
    for paragraph in paragraphs:
        if len(paragraph) < 180:
            continue
        key = re.sub(r"\s+", " ", paragraph.lower())
        key = re.sub(r"\(s\d+(?:,\s*s\d+)*\)", "(citation)", key)
        normalized.append(key)
        original.setdefault(key, paragraph)
    counts = Counter(normalized)
    return [(original[key], count) for key, count in counts.items() if count > 1]


def first_index(paragraphs: list[str], needles: list[str]) -> int | None:
    lowered = [p.lower() for p in paragraphs]
    for index, paragraph in enumerate(lowered):
        if any(needle.lower() == paragraph for needle in needles):
            return index
    return None


def score_rows_have_rationale(doc, heading_blob: str) -> bool:
    if "bewertungslogik" in heading_blob or "scoring logic" in heading_blob:
        return True
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        joined = " | ".join(headers)
        if "technologie" in joined or "verfahren" in joined:
            return any(
                "evidenz" in header or "rationale" in header or "begründ" in header or "begruend" in header
                for header in headers
            )
    return False


def emit(args, failures, warnings, doc=None, table_roles=None) -> None:
    payload = {
        "ok": not failures,
        "docx": str(args.docx),
        "failures": failures,
        "warnings": warnings,
    }
    if doc is not None:
        payload["metrics"] = {
            "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "tables": len(doc.tables),
            "images": len(doc.inline_shapes),
            "domain_table_roles": sorted(table_roles or []),
        }
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    raise SystemExit(0 if not failures else 2)


if __name__ == "__main__":
    main()
