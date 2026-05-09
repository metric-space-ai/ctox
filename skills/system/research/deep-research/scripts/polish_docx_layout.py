#!/usr/bin/env python3
"""Post-process a rendered deep-research DOCX for client-facing layout.

The manuscript renderer is intentionally conservative: it creates a
semantically correct Word file from structured report data. This helper is
the second pass for visual quality. It is generic across report topics and
must not rely on domain-specific wording.

Typical use:

    python3 polish_docx_layout.py --in report.raw.docx --out report.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


def _missing_python_docx() -> None:
    sys.stderr.write(
        "python-docx not installed. Run: python3 -m pip install python-docx\n"
    )
    sys.exit(2)


try:
    from docx import Document
    from docx.enum.table import (
        WD_CELL_VERTICAL_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_TABLE_ALIGNMENT,
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    _missing_python_docx()
    raise


BLUE = "1F4E79"
BLUE_DARK = "17365D"
GREY = "666666"
GRID = "B7C7D6"
ROW_ALT = "F3F7FA"
HEADER_RULE = "D9E2EC"


def remove_paragraph(paragraph) -> None:
    node = paragraph._element
    node.getparent().remove(node)
    paragraph._p = paragraph._element = None


def has_page_break(paragraph) -> bool:
    xml = paragraph._p.xml
    return "<w:br" in xml and 'w:type="page"' in xml


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def paragraph_bottom_border(paragraph, color=BLUE, size="10", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)


def style_run(run, size=None, bold=None, italic=None, color=None, font="Arial") -> None:
    run.font.name = font
    rpr = run._element.rPr
    if rpr is not None:
        rpr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def first_report_title(doc) -> str:
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            return text
    return ""


def short_header_title(title: str) -> str:
    title = title.strip()
    if not title:
        return "Machbarkeitsstudie"
    lowered = title.lower()
    prefix = "Machbarkeitsstudie"
    if "machbarkeitsstudie" not in lowered and "feasibility" not in lowered:
        prefix = "Report"
    cleaned = title
    if cleaned.lower().startswith(prefix.lower()):
        cleaned = cleaned[len(prefix) :].strip(" :-")
    if len(cleaned) > 78:
        cleaned = cleaned[:75].rstrip() + "..."
    return f"{prefix} | {cleaned}" if cleaned else prefix


def remove_placeholder_toc_block(doc) -> int:
    """Remove only the generated placeholder TOC block, never a real TOC."""
    paragraphs = list(doc.paragraphs)
    start = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().lower() in {"inhaltsverzeichnis", "table of contents"}:
            start = i
            break
    if start is None:
        return 0

    lookahead = "\n".join(p.text for p in paragraphs[start : min(start + 8, len(paragraphs))])
    placeholder_markers = (
        "rechtsklick",
        "field",
        "aktualisieren",
        "update",
        "word fills",
    )
    if not any(marker in lookahead.lower() for marker in placeholder_markers):
        return 0

    end = start
    for j in range(start + 1, len(paragraphs)):
        end = j
        if has_page_break(paragraphs[j]):
            break
    for paragraph in paragraphs[start : end + 1]:
        remove_paragraph(paragraph)
    return end - start + 1


def add_headers_footers(doc, header_text: str) -> None:
    for idx, section in enumerate(doc.sections):
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.15)
        section.right_margin = Cm(2.15)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        if idx > 0:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.text = header_text
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in header.runs:
            style_run(run, size=8, color=GREY)
        paragraph_bottom_border(header, color=HEADER_RULE, size="4", space="2")

        footer = section.footer.paragraphs[0]
        footer.text = ""
        add_page_number(footer)


def polish_cover(doc) -> None:
    if not doc.paragraphs:
        return
    title = doc.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(56)
    title.paragraph_format.space_after = Pt(12)
    paragraph_bottom_border(title, color=BLUE, size="14", space="10")
    for run in title.runs:
        style_run(run, size=21, bold=True, color=BLUE_DARK)

    if len(doc.paragraphs) > 1:
        subtitle = doc.paragraphs[1]
        subtitle.paragraph_format.space_before = Pt(6)
        subtitle.paragraph_format.space_after = Pt(12)
        for run in subtitle.runs:
            style_run(run, size=11.5, italic=True, color=GREY)

    if len(doc.paragraphs) > 2:
        version = doc.paragraphs[2]
        version.paragraph_format.space_after = Pt(36)
        for run in version.runs:
            style_run(run, size=9.5, color=GREY)

    for paragraph in doc.paragraphs[:8]:
        if "nur" in paragraph.text.lower() and "beratung" in paragraph.text.lower():
            paragraph.paragraph_format.left_indent = Cm(0.15)
            paragraph.paragraph_format.right_indent = Cm(0.15)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                style_run(run, size=9.8, italic=True, color=GREY)
            break


def apply_paragraph_style(paragraph) -> None:
    name = paragraph.style.name
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.08

    if name == "Heading 1":
        fmt.space_before = Pt(20)
        fmt.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_bottom_border(paragraph, BLUE, size="8", space="5")
        for run in paragraph.runs:
            style_run(run, size=16, bold=True, color=BLUE_DARK)
    elif name == "Heading 2":
        fmt.space_before = Pt(13)
        fmt.space_after = Pt(4)
        for run in paragraph.runs:
            style_run(run, size=12.5, bold=True, color=BLUE_DARK)
    elif paragraph.text.startswith(("Tabelle ", "Table ", "Abbildung ", "Figure ")):
        fmt.space_before = Pt(5)
        fmt.space_after = Pt(4)
        for run in paragraph.runs:
            style_run(run, size=8.5, italic=True, color=GREY)
    else:
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(5)
        for run in paragraph.runs:
            style_run(run, size=10.3)


def wide_method_matrix_headers(headers: list[str]) -> bool:
    normalized = [h.strip().lower() for h in headers]
    required = {"methode", "kopplung", "zielinformation"}
    if not required.issubset(set(normalized)):
        return False
    return len(headers) >= 7 and any(h in normalized for h in ("grenzen", "limitations"))


def reflow_wide_method_matrices(doc) -> int:
    """Convert a cramped method-screening table into four readable columns.

    The trigger is structural, not topic-specific: a wide matrix with method,
    coupling, target information, strengths, limits, and assessment columns.
    """
    changed = 0
    for table in list(doc.tables):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if not wide_method_matrix_headers(headers):
            continue

        header_lut = {h.strip().lower(): idx for idx, h in enumerate(headers)}
        idx_method = header_lut.get("methode", 0)
        idx_coupling = header_lut.get("kopplung", 1)
        idx_target = header_lut.get("zielinformation", 2)
        idx_strengths = header_lut.get("staerken", header_lut.get("stärken", 3))
        idx_limits = header_lut.get("grenzen", header_lut.get("limitations", 4))

        rows: list[list[str]] = []
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            method = cells[idx_method] if idx_method < len(cells) else ""
            coupling = cells[idx_coupling] if idx_coupling < len(cells) else ""
            target = cells[idx_target] if idx_target < len(cells) else ""
            strengths = cells[idx_strengths] if idx_strengths < len(cells) else ""
            limits = cells[idx_limits] if idx_limits < len(cells) else ""
            profile = "; ".join(part for part in (coupling, target) if part)

            assessment_parts = [limits] if limits else []
            for idx, header in enumerate(headers):
                if idx in {idx_method, idx_coupling, idx_target, idx_strengths, idx_limits}:
                    continue
                value = cells[idx] if idx < len(cells) else ""
                if value:
                    assessment_parts.append(f"{header}: {value}")
            rows.append([method, profile, strengths, " ".join(assessment_parts)])

        if not rows:
            continue

        new_table = doc.add_table(rows=1, cols=4)
        for idx, text in enumerate(["Methode", "Einsatzprofil", "Stärken", "Grenzen / Bewertung"]):
            new_table.rows[0].cells[idx].text = text
        for values in rows:
            row_cells = new_table.add_row().cells
            for idx, text in enumerate(values):
                row_cells[idx].text = text

        table._tbl.addprevious(new_table._tbl)
        table._tbl.getparent().remove(table._tbl)
        changed += 1
    return changed


def insert_page_breaks_for_large_tables(doc) -> int:
    """Nudge large tables away from the bottom of a page without domain words."""
    changed = 0
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text.startswith(("Tabelle ", "Table ")):
            continue
        # Large matrices are easier to review when introduced from a fresh page.
        if any(word in text.lower() for word in ("szenario", "scenario", "matrix")):
            paragraph.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
            changed += 1
    return changed


def polish_tables(doc) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        if table.rows:
            set_repeat_table_header(table.rows[0])

        col_count = len(table.columns)
        if col_count == 3:
            widths = [2100, 3000, 3600]
        elif col_count == 4:
            widths = [1700, 2500, 2500, 3100]
        elif col_count == 5:
            widths = [1650, 2000, 2200, 1900, 950]
        else:
            widths = [int(8700 / max(col_count, 1))] * col_count

        for r_idx, row in enumerate(table.rows):
            prevent_row_split(row)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(18)
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=85, start=95, bottom=85, end=95)
                if c_idx < len(widths):
                    set_width(cell, widths[c_idx])
                if r_idx == 0:
                    header_text = cell.text.strip()
                    compact_headers = {
                        "Mess- oder Simulationsmethode": "Mess-/Sim.-Methode",
                        "Mess-/Simulationsmethode": "Mess-/Sim.-Methode",
                        "Measurement or simulation method": "Measurement/simulation method",
                    }
                    if header_text in compact_headers:
                        cell.text = compact_headers[header_text]
                    set_cell_shading(cell, BLUE)
                elif r_idx % 2 == 1:
                    set_cell_shading(cell, ROW_ALT)
                else:
                    set_cell_shading(cell, "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        if r_idx == 0:
                            style_run(run, size=8.6, bold=True, color="FFFFFF")
                        else:
                            style_run(run, size=8.3)


def compact_reference_appendix(doc) -> None:
    in_sources = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if (
            text.startswith("anhang")
            and any(marker in text for marker in ("quelle", "reference", "evidence"))
        ):
            in_sources = True
            continue
        if not in_sources:
            continue
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            style_run(run, size=8.35)


def polish_docx(input_path: Path, output_path: Path) -> dict[str, int]:
    doc = Document(input_path)
    title = first_report_title(doc)
    stats = {
        "placeholder_toc_paragraphs_removed": remove_placeholder_toc_block(doc),
        "wide_matrices_reflowed": reflow_wide_method_matrices(doc),
        "large_table_page_breaks": insert_page_breaks_for_large_tables(doc),
    }
    add_headers_footers(doc, short_header_title(title))
    polish_cover(doc)
    for paragraph in doc.paragraphs:
        apply_paragraph_style(paragraph)
    compact_reference_appendix(doc)
    polish_tables(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    stats["bytes"] = output_path.stat().st_size
    return stats


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--in", dest="input_path", required=True, type=Path)
    parser.add_argument("--out", dest="output_path", required=True, type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.input_path.exists():
        sys.stderr.write(f"input DOCX not found: {args.input_path}\n")
        return 1
    stats = polish_docx(args.input_path, args.output_path)
    print(
        "OK {bytes} {path} toc_removed={toc} wide_matrices={wide}".format(
            bytes=stats["bytes"],
            path=args.output_path,
            toc=stats["placeholder_toc_paragraphs_removed"],
            wide=stats["wide_matrices_reflowed"],
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
