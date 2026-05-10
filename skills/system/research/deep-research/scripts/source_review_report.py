#!/usr/bin/env python3
"""Build a client-facing DOCX/XLSX report from source-review artifacts."""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any


def load_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def clean(value: Any, max_len: int | None = None) -> str:
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", " ", str(value or ""))
    text = re.sub(r"\s+", " ", text).strip()
    if max_len and len(text) > max_len:
        return text[: max_len - 1].rstrip() + "..."
    return text


def int_or_zero(value: Any) -> int:
    try:
        return int(float(str(value or "0").replace(",", "")))
    except Exception:
        return 0


def normalize_doi(value: str) -> str:
    value = clean(value)
    value = re.sub(r"^https?://(dx\.)?doi\.org/", "", value, flags=re.IGNORECASE)
    return value


def source_key(row: dict[str, str]) -> str:
    return row.get("openalex_id") or normalize_doi(row.get("doi") or row.get("url", "")) or row.get("url", "") or row.get("title", "")


def status_index(reading_rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    out: dict[str, dict[str, str]] = {}
    for row in reading_rows:
        keys = [
            row.get("openalex_id", ""),
            normalize_doi(row.get("doi", "")),
            row.get("read_url", ""),
            row.get("title", ""),
        ]
        for key in keys:
            if key:
                out[key] = row
    return out


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    # Keep reports readable even when the runtime lacks python-docx internals.
    run = paragraph.add_run(text)
    run.font.color.rgb = paragraph.part.document.styles["Hyperlink"].font.color.rgb if "Hyperlink" in paragraph.part.document.styles else None
    paragraph.add_run(f" ({url})")


def add_table(document: Any, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = clean(value)
    document.add_paragraph()


def infer_mass_limit_kg(topic: str) -> float | None:
    patterns = [
        r"(?:up to|under|below|<=|≤|max(?:imum)?)\s*(\d+(?:\.\d+)?)\s*kg",
        r"(\d+(?:\.\d+)?)\s*kg\s*(?:or less|and below|and under)",
    ]
    lowered = topic.lower()
    for pattern in patterns:
        match = re.search(pattern, lowered)
        if match:
            try:
                return float(match.group(1))
            except Exception:
                return None
    return None


def mass_values_kg(row: dict[str, str]) -> list[float]:
    if row.get("family") != "mass_payload":
        return []
    unit = clean(row.get("unit")).lower()
    if unit not in {"kg", "g", "lb", "lbs"}:
        return []
    values = [float(item) for item in re.findall(r"-?\d+(?:\.\d+)?", clean(row.get("value")).replace(",", ""))]
    if unit == "g":
        return [value / 1000 for value in values]
    if unit in {"lb", "lbs"}:
        return [value * 0.45359237 for value in values]
    return values


def measurement_scope(row: dict[str, str], mass_limit_kg: float | None) -> str:
    masses = mass_values_kg(row)
    if not masses or mass_limit_kg is None:
        return "not_mass_scoped"
    if any(value < 0 for value in masses):
        return "invalid_or_context"
    if any(value > mass_limit_kg for value in masses):
        return "outside_mass_scope"
    return "inside_mass_scope"


def set_doc_style(document: Any) -> None:
    from docx.shared import Pt

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 26), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        if style_name in styles:
            styles[style_name].font.name = "Aptos"
            styles[style_name].font.size = Pt(size)


def add_bullets(document: Any, bullets: list[str]) -> None:
    for item in bullets:
        document.add_paragraph(item, style="List Bullet")


def build_interpretation(topic: str, candidate_rows: list[dict[str, str]], measurement_rows: list[dict[str, str]], off_scope_rows: list[dict[str, str]]) -> list[str]:
    families = Counter(row.get("family", "unknown") for row in measurement_rows)
    accepted = len(candidate_rows)
    bullets = [
        f"The accepted source catalog contains {accepted} topic-relevant sources after deterministic screening.",
        "The reading stage separates full/readable evidence from metadata-only or blocked sources, so the report does not overclaim access.",
    ]
    if families:
        strongest = ", ".join(f"{family} ({count})" for family, count in families.most_common(4))
        bullets.append(f"The strongest extracted measurement families in the current reading pass are: {strongest}.")
    if off_scope_rows:
        bullets.append(
            f"{len(off_scope_rows)} extracted mass rows were kept out of the main evidence table because they are outside the requested mass scope or context-only."
        )
    lowered = topic.lower()
    if any(term in lowered for term in ["drone", "uav", "uas", "unmanned"]):
        bullets.extend(
            [
                "For small UAVs, the strongest public evidence is propulsion, wind-tunnel and propeller test data; direct structural load datasets are materially scarcer.",
                "Payload and takeoff-weight evidence is often available through specifications, classification papers and application studies, but it is less standardized than thrust/propulsion measurements.",
                "A defensible final research deliverable should keep three buckets separate: source discovery, readable source evidence, and extracted numeric measurements.",
            ]
        )
    return bullets


def build_docx(
    title: str,
    topic: str,
    discovery_dir: Path,
    reading_dir: Path,
    out_path: Path,
) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise SystemExit("python-docx is required. Use the bundled Codex workspace Python runtime.") from exc

    candidates = load_csv(discovery_dir / "candidate_sources.csv")
    screened = load_csv(discovery_dir / "screened_sources.csv")
    rejected = load_csv(discovery_dir / "rejected_sources.csv")
    search_protocol = load_csv(discovery_dir / "search_protocol.csv")
    discovery_graph = load_json(discovery_dir / "discovery_graph.json")
    reading_rows = load_csv(reading_dir / "reading_status.csv")
    measurements = load_csv(reading_dir / "extracted_measurements.csv")
    discovery_summary = load_json(discovery_dir / "summary.json")
    reading_summary = load_json(reading_dir / "reading_summary.json")
    mass_limit_kg = infer_mass_limit_kg(topic)
    scoped_measurements = [
        row for row in measurements if measurement_scope(row, mass_limit_kg) not in {"outside_mass_scope", "invalid_or_context"}
    ]
    off_scope_measurements = [
        row for row in measurements if measurement_scope(row, mass_limit_kg) in {"outside_mass_scope", "invalid_or_context"}
    ]

    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.add_heading(title or "Source Review Report", 0)
    doc.add_paragraph(f"Scope: {topic}")
    doc.add_paragraph("Generated from deterministic discovery, reading and extraction artifacts.")

    doc.add_heading("Executive Summary", level=1)
    accepted_count = len(candidates)
    reviewed_result_records = int_or_zero(discovery_summary.get("reviewed_results"))
    screened_unique_count = int_or_zero(discovery_summary.get("screened_unique_sources")) or len(screened) or len(candidates) + len(rejected)
    rejected_count = len(rejected)
    readable = int_or_zero(reading_summary.get("readable_sources"))
    metadata_only = int_or_zero(reading_summary.get("metadata_only_sources"))
    blocked = int_or_zero(reading_summary.get("blocked_sources"))
    extracted_sources = int_or_zero(reading_summary.get("extracted_sources"))
    measurement_count = len(scoped_measurements)
    add_bullets(
        doc,
        [
            f"Reviewed result records: {reviewed_result_records or screened_unique_count}. Unique screened sources: {screened_unique_count}. Accepted relevant sources: {accepted_count}. Rejected/off-topic records: {rejected_count}.",
            f"Targeted reading pass: {readable} readable sources, {metadata_only} metadata-only sources and {blocked} blocked sources.",
            f"In-scope extracted measurement evidence: {measurement_count} evidence rows from {extracted_sources} sources.",
            "The report distinguishes discovery coverage from source readability and extracted evidence; inaccessible sources are not treated as reviewed full text.",
        ],
    )

    doc.add_heading("Research Method", level=1)
    doc.add_paragraph(
        "The workflow used three controlled stages: broad source discovery, targeted reading/source resolution, and evidence extraction. "
        "Discovery screened search results and metadata for topic relevance. Reading then resolved open-access locations and source pages. "
        "Extraction captured numeric measurement evidence with source snippets and normalized measurement families."
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Search/query paths", str(len(search_protocol))],
            ["Reviewed result records", str(reviewed_result_records or "not logged")],
            ["Unique screened sources", str(screened_unique_count)],
            ["Accepted sources", str(accepted_count)],
            ["Rejected records", str(rejected_count)],
            ["Discovery graph nodes", str(len(discovery_graph.get("nodes", []))) if discovery_graph else "not generated"],
            ["Discovery graph edges", str(len(discovery_graph.get("edges", []))) if discovery_graph else "not generated"],
            ["Readable sources", str(readable)],
            ["Metadata-only sources", str(metadata_only)],
            ["Blocked sources", str(blocked)],
            ["In-scope measurement rows", str(measurement_count)],
            ["Off-scope/context measurement rows excluded", str(len(off_scope_measurements))],
        ],
    )

    doc.add_heading("Source Landscape", level=1)
    reasons = Counter(clean(row.get("acceptance_reason")) for row in candidates)
    statuses = Counter(row.get("status", "unknown") for row in reading_rows)
    add_table(
        doc,
        ["Accepted source gate", "Count"],
        [[reason or "unspecified", str(count)] for reason, count in reasons.most_common(12)],
    )
    add_table(
        doc,
        ["Reading status", "Count", "Meaning"],
        [
            ["extracted", str(statuses.get("extracted", 0)), "Readable and numeric evidence extracted"],
            ["readable_no_measurements", str(statuses.get("readable_no_measurements", 0)), "Readable but no numeric evidence found in current pass"],
            ["metadata_only", str(statuses.get("metadata_only", 0)), "Only abstract/metadata used"],
            ["blocked", str(statuses.get("blocked", 0)), "No readable source text available from attempted URLs"],
        ],
    )

    doc.add_heading("Findings", level=1)
    add_bullets(doc, build_interpretation(topic, candidates, scoped_measurements, off_scope_measurements))

    doc.add_heading("Extracted Measurement Evidence", level=1)
    if scoped_measurements:
        family_counts = Counter(row.get("family", "unknown") for row in scoped_measurements)
        add_table(doc, ["Measurement family", "Evidence rows"], [[k, str(v)] for k, v in family_counts.most_common()])
        measurement_table = []
        for row in scoped_measurements[:35]:
            measurement_table.append(
                [
                    clean(row.get("family"), 32),
                    clean(f"{row.get('value', '')} {row.get('unit', '')}", 24),
                    clean(row.get("title"), 74),
                    clean(row.get("snippet"), 180),
                ]
            )
        add_table(doc, ["Family", "Value", "Source", "Evidence snippet"], measurement_table)
    else:
        doc.add_paragraph("No measurement rows were extracted in the current reading pass.")

    if off_scope_measurements:
        doc.add_heading("Excluded Measurement Rows", level=2)
        doc.add_paragraph(
            "These rows were extracted by the reader but kept out of the main evidence table because they are outside the requested mass scope or context-only."
        )
        add_table(
            doc,
            ["Scope reason", "Value", "Source"],
            [
                [
                    measurement_scope(row, mass_limit_kg),
                    clean(f"{row.get('value', '')} {row.get('unit', '')}", 24),
                    clean(row.get("title"), 95),
                ]
                for row in off_scope_measurements[:20]
            ],
        )

    doc.add_heading("Priority Source Catalog", level=1)
    index = status_index(reading_rows)
    top_sources: list[list[str]] = []
    for row in sorted(candidates, key=lambda item: int_or_zero(item.get("relevance_score")), reverse=True)[:40]:
        status = index.get(row.get("openalex_id", "")) or index.get(normalize_doi(row.get("doi", ""))) or {}
        top_sources.append(
            [
                clean(row.get("relevance_score"), 8),
                clean(status.get("status", "not_read_in_pass"), 24),
                clean(row.get("title"), 88),
                clean(normalize_doi(row.get("doi") or row.get("url", "")), 42),
                clean(row.get("acceptance_reason"), 36),
            ]
        )
    add_table(doc, ["Score", "Reading status", "Title", "DOI/URL", "Gate"], top_sources)

    doc.add_heading("Coverage Gaps and Next Work", level=1)
    gaps = [
        "Expand the reading pass for source families that remained metadata-only or blocked.",
        "Separate manufacturer specification data, regulatory/classification data, propulsion bench data and flight-test logs in the final evidence model.",
        "For load-data questions, add a second extraction pass that targets tables, supplementary files and dataset repositories explicitly.",
        "For paywalled DOI-heavy clusters, record metadata coverage separately from reviewed full text and seek legally accessible alternate copies.",
    ]
    add_bullets(doc, gaps)

    doc.add_heading("Artifact Appendix", level=1)
    add_table(
        doc,
        ["Companion artifact", "File"],
        [
            ["Accepted source catalog", "candidate_sources.csv"],
            ["Rejected/off-topic audit", "rejected_sources.csv"],
            ["Discovery graph", "discovery_graph.json"],
            ["Reading status ledger", "reading_status.csv"],
            ["Extracted measurements", "extracted_measurements.csv"],
            ["Reading graph", "reading_graph.json"],
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return {
        "docx": str(out_path),
        "accepted_sources": accepted_count,
        "reviewed_result_records": reviewed_result_records,
        "screened_unique_sources": screened_unique_count,
        "readable_sources": readable,
        "metadata_only_sources": metadata_only,
        "blocked_sources": blocked,
        "in_scope_measurement_rows": measurement_count,
        "off_scope_measurement_rows": len(off_scope_measurements),
    }


def autosize_sheet(sheet: Any) -> None:
    for column in sheet.columns:
        width = 10
        col_letter = column[0].column_letter
        for cell in column:
            width = min(70, max(width, len(str(cell.value or "")) + 2))
        sheet.column_dimensions[col_letter].width = width


def build_xlsx(discovery_dir: Path, reading_dir: Path, out_path: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except Exception:
        return

    wb = Workbook()
    default = wb.active
    wb.remove(default)
    datasets = [
        ("accepted_sources", load_csv(discovery_dir / "candidate_sources.csv")),
        ("reading_status", load_csv(reading_dir / "reading_status.csv")),
        ("measurements", load_csv(reading_dir / "extracted_measurements.csv")),
        ("rejected_sources", load_csv(discovery_dir / "rejected_sources.csv")),
    ]
    for name, rows in datasets:
        sheet = wb.create_sheet(name[:31])
        if not rows:
            sheet.append(["empty"])
            continue
        headers = list(rows[0].keys())
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for row in rows:
            sheet.append([clean(row.get(header, ""), 32000) for header in headers])
        sheet.freeze_panes = "A2"
        autosize_sheet(sheet)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--topic", required=True)
    parser.add_argument("--title", default="Source Review Report")
    parser.add_argument("--discovery-dir", type=Path, required=True)
    parser.add_argument("--reading-dir", type=Path, required=True)
    parser.add_argument("--out-docx", type=Path, required=True)
    parser.add_argument("--out-xlsx", type=Path)
    args = parser.parse_args(argv)

    summary = build_docx(args.title, args.topic, args.discovery_dir, args.reading_dir, args.out_docx)
    if args.out_xlsx:
        build_xlsx(args.discovery_dir, args.reading_dir, args.out_xlsx)
        summary["xlsx"] = str(args.out_xlsx)
    print(json.dumps(summary, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
