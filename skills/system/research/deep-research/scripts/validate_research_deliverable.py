#!/usr/bin/env python3
"""Validate a CTOX deep-research workspace and final DOCX deliverable."""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path
from typing import Any


PLACEHOLDER_MARKERS = [
    "status: provisional",
    "placeholder",
    "research underway",
    "checks planned",
    "to be completed",
    "tbd",
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workspace", required=True, type=Path)
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--min-sources", type=int, default=20)
    parser.add_argument("--min-reads", type=int, default=5)
    parser.add_argument("--min-draft-chars", type=int, default=8000)
    parser.add_argument("--require-call-counts", action="store_true")
    args = parser.parse_args()

    failures: list[str] = []
    warnings: list[str] = []

    workspace = args.workspace
    docx = args.docx

    if not workspace.is_dir():
        failures.append(f"research workspace is missing or not a directory: {workspace}")
    else:
        validate_workspace(workspace, args, failures, warnings)

    validate_docx(docx, failures, warnings)

    result = {
        "ok": not failures,
        "workspace": str(workspace),
        "docx": str(docx),
        "failures": failures,
        "warnings": warnings,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if not failures else 2)


def validate_workspace(
    workspace: Path,
    args: argparse.Namespace,
    failures: list[str],
    warnings: list[str],
) -> None:
    dot_docx_dirs = [path for path in workspace.parent.glob("*.docx") if path.is_dir()]
    if dot_docx_dirs:
        failures.append(
            "found directory with .docx suffix; final Word output must be a file: "
            + ", ".join(str(path) for path in dot_docx_dirs)
        )

    required = [
        "manifest.json",
        "evidence_bundle.json",
        "sources.jsonl",
        "data_links.json",
        "CONTINUE.md",
        "synthesis/evidence-matrix.md",
        "synthesis/technology-scores.md",
        "synthesis/report-outline.md",
        "synthesis/figure-plan.md",
        "synthesis/report-draft.md",
        "synthesis/qa-notes.md",
    ]
    for relative in required:
        path = workspace / relative
        if not path.is_file():
            failures.append(f"missing required workspace file: {path}")

    sources_path = workspace / "sources.jsonl"
    source_count = count_jsonl(sources_path)
    if source_count < args.min_sources:
        failures.append(
            f"sources.jsonl has {source_count} source(s); expected at least {args.min_sources}"
        )

    read_count = len([path for path in (workspace / "reads").glob("*") if path.is_file()])
    if read_count < args.min_reads:
        failures.append(f"reads/ has {read_count} file(s); expected at least {args.min_reads}")

    evidence_path = workspace / "evidence_bundle.json"
    evidence = read_json(evidence_path, failures)
    if isinstance(evidence, dict):
        evidence_sources = extract_sources(evidence)
        if len(evidence_sources) < args.min_sources:
            failures.append(
                f"evidence_bundle.json has {len(evidence_sources)} source(s); expected at least {args.min_sources}"
            )
        call_counts = evidence.get("research_call_counts") or evidence.get("call_counts")
        if args.require_call_counts and not call_counts:
            failures.append("missing research_call_counts/call_counts in evidence_bundle.json")
        if call_counts:
            warnings.append(f"call_counts={call_counts}")

    draft_path = workspace / "synthesis" / "report-draft.md"
    draft_text = read_text(draft_path)
    if len(draft_text.strip()) < args.min_draft_chars:
        failures.append(
            f"report-draft.md has {len(draft_text.strip())} chars; expected at least {args.min_draft_chars}"
        )
    check_placeholder_markers(draft_path, draft_text, failures)

    for relative in [
        "synthesis/evidence-matrix.md",
        "synthesis/technology-scores.md",
        "synthesis/report-outline.md",
        "synthesis/figure-plan.md",
        "synthesis/qa-notes.md",
    ]:
        path = workspace / relative
        text = read_text(path)
        if len(text.strip()) < 500 and relative not in {"synthesis/figure-plan.md", "synthesis/qa-notes.md"}:
            failures.append(f"{relative} is too short to be decision-grade ({len(text.strip())} chars)")
        check_placeholder_markers(path, text, failures)


def validate_docx(docx: Path, failures: list[str], warnings: list[str]) -> None:
    if docx.is_dir():
        failures.append(f"final DOCX path is a directory, not a file: {docx}")
        return
    if not docx.is_file():
        failures.append(f"final DOCX is missing: {docx}")
        return
    if docx.suffix.lower() != ".docx":
        failures.append(f"final document does not have .docx suffix: {docx}")
    if docx.stat().st_size < 20_000:
        failures.append(f"final DOCX is suspiciously small ({docx.stat().st_size} bytes)")
    try:
        with zipfile.ZipFile(docx) as archive:
            names = set(archive.namelist())
            for required in {"[Content_Types].xml", "word/document.xml"}:
                if required not in names:
                    failures.append(f"DOCX zip missing required member: {required}")
    except zipfile.BadZipFile:
        failures.append(f"final DOCX is not a valid ZIP/DOCX file: {docx}")
        return

    try:
        from docx import Document  # type: ignore

        document = Document(str(docx))
        paragraph_count = len(document.paragraphs)
        table_count = len(document.tables)
        image_count = len(document.inline_shapes)
        warnings.append(
            f"docx_structure={{paragraphs:{paragraph_count},tables:{table_count},images:{image_count}}}"
        )
        if paragraph_count < 40:
            failures.append(f"DOCX has too few paragraphs for a feasibility study: {paragraph_count}")
        if table_count < 2:
            failures.append(f"DOCX has too few tables for a feasibility study: {table_count}")
    except Exception as exc:  # pragma: no cover - depends on optional python-docx
        warnings.append(f"python-docx structural check skipped/failed: {exc}")


def count_jsonl(path: Path) -> int:
    if not path.is_file():
        return 0
    return sum(1 for line in path.read_text(encoding="utf-8", errors="replace").splitlines() if line.strip())


def read_json(path: Path, failures: list[str]) -> Any:
    if not path.is_file():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        failures.append(f"invalid JSON in {path}: {exc}")
        return None


def read_text(path: Path) -> str:
    if not path.is_file():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def extract_sources(payload: dict[str, Any]) -> list[Any]:
    for key in ["sources", "results", "papers", "evidence"]:
        value = payload.get(key)
        if isinstance(value, list):
            return value
    return []


def check_placeholder_markers(path: Path, text: str, failures: list[str]) -> None:
    lowered = text.lower()
    found = [marker for marker in PLACEHOLDER_MARKERS if marker in lowered]
    if found:
        failures.append(f"{path} contains placeholder marker(s): {', '.join(found)}")


if __name__ == "__main__":
    main()
